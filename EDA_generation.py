import pandas as pd
import numpy as np
import seaborn as sns
import matplotlib.pyplot as plt
from docx import Document
import great_expectations as ge


file_path = r"C:\Users\kkewa\TATA-GenAI-Powered-Data-Analytics-job-sim\Delinquency_prediction_dataset.xlsx"  

df = pd.read_excel(file_path)


num_records = df.shape[0]
num_columns = df.shape[1]
columns = df.columns.tolist()

data_types = df.dtypes.apply(lambda x: str(x)).to_dict()


missing = df.isnull().sum()
missing_percent = (missing / len(df)) * 100

missing_df = pd.DataFrame({
    "column": df.columns,
    "missing_count": missing.values,
    "missing_percent": missing_percent.values
})

missing_df = missing_df[missing_df["missing_count"] > 0]

duplicates = df.duplicated().sum()


numeric_df = df.select_dtypes(include=[np.number])
summary_stats = numeric_df.describe()


correlation_matrix = numeric_df.corr()


plt.figure(figsize=(10,8))
sns.heatmap(correlation_matrix, annot=False, cmap="coolwarm")
plt.title("Correlation Heatmap")
plt.savefig("correlation.png")
plt.close()

anomalies = {}

for col in numeric_df.columns:
    q1 = numeric_df[col].quantile(0.25)
    q3 = numeric_df[col].quantile(0.75)
    iqr = q3 - q1
    outliers = numeric_df[(numeric_df[col] < (q1 - 1.5 * iqr)) |
                          (numeric_df[col] > (q3 + 1.5 * iqr))]
    anomalies[col] = len(outliers)

risk_indicators = []

for col in correlation_matrix.columns:
    if "delinq" in col.lower() or "default" in col.lower():
        target_col = col
        break
else:
    target_col = None

if target_col:
    corr_with_target = correlation_matrix[target_col].sort_values(ascending=False)
    top_features = corr_with_target[1:4]  # top 3 excluding itself

    for feature, value in top_features.items():
        risk_indicators.append(f"{feature} (correlation: {round(value,2)})")


doc = Document("EDA_SummaryReport_Template.docx")


doc.paragraphs[0].add_run(
    "\nThis report analyzes Geldium's dataset to assess data quality, identify missing values, "
    "detect anomalies, and uncover early indicators of credit delinquency risk."
)


doc.add_paragraph(f"\nNumber of records: {num_records}")
doc.add_paragraph(f"Number of columns: {num_columns}")
doc.add_paragraph(f"Duplicate records: {duplicates}")
doc.add_paragraph(f"Columns: {', '.join(columns[:10])}...")

doc.add_paragraph("\nMissing Data Summary:")

for _, row in missing_df.iterrows():
    doc.add_paragraph(
        f"{row['column']} → {int(row['missing_count'])} missing "
        f"({round(row['missing_percent'],2)}%)"
    )

doc.add_paragraph("\nKey Findings & Risk Indicators:")

if risk_indicators:
    for r in risk_indicators:
        doc.add_paragraph(f"- {r}")
else:
    doc.add_paragraph("Target variable not clearly identified for correlation analysis.")

doc.add_paragraph("\nAnomalies (Outliers Count):")
for col, count in anomalies.items():
    if count > 0:
        doc.add_paragraph(f"{col}: {count} outliers")


doc.add_paragraph(
    "\nGenAI tools were used to assist in identifying patterns, suggesting imputation strategies, "
    "and highlighting potential risk indicators while ensuring no sensitive financial data was exposed."
)


doc.add_paragraph(
    "\nThe dataset shows areas of missing data and potential anomalies that must be addressed before modeling. "
    "Next steps include data cleaning, feature engineering, and building predictive models for delinquency risk."
)


doc.save("EDA_Final_Report.docx")

print("✅ Report generated: EDA_Final_Report.docx")